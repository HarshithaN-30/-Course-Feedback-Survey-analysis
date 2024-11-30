from flask import Flask, request, render_template, send_file
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Correct import for paragraph alignment
from docx.shared import RGBColor  # Correct import for RGB color
import os

# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['OUTPUT_FOLDER'] = 'output/'

# Create necessary folders
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

if not os.path.exists(app.config['OUTPUT_FOLDER']):
    os.makedirs(app.config['OUTPUT_FOLDER'])

# Home route
@app.route('/')
def upload_page():
    return render_template('upload.html', download_link=None)

# Process uploaded file
@app.route('/process', methods=['POST'])
def process_file():
    if 'file' not in request.files:
        return render_template('upload.html', error="No file uploaded!", download_link=None)

    file = request.files['file']
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    # Process the Excel file
    output_path = process_excel(file_path)
    download_link = f"/output/{os.path.basename(output_path)}"

    # Render the same page with the download link
    return render_template('upload.html', download_link=download_link)

# Serve output files for download
@app.route('/output/<filename>')
def download_file(filename):
    return send_file(os.path.join(app.config['OUTPUT_FOLDER'], filename), as_attachment=True)

# Process Excel file and generate Word report
def process_excel(file_path):
    # Load the Excel file
    df = pd.read_excel(file_path, header=None)

    # Check if metadata exists
    metadata = extract_metadata(df)

    if metadata:  # Case 1: Metadata exists
        question_start_row = find_question_start(df)  # Find survey start row
        df_survey = pd.read_excel(file_path, skiprows=question_start_row)
    else:  # Case 2: No metadata
        df_survey = pd.read_excel(file_path)  # Read directly, assume first row is header
        metadata = {"Subject Code": "Unknown", "Subject Name": "Unknown", "Branch": "Unknown", "Year": "Unknown"}

    # Drop unnecessary columns: NAME and USN
    df_survey.drop(columns=['NAME', 'USN'], errors='ignore', inplace=True)

    # Get the course code from metadata
    course_code = metadata.get('Subject Code', 'Unknown')

    # Rename survey columns based on the course code
    for idx, col in enumerate(df_survey.columns[1:], start=1):
        course_code_column_name = f"{course_code}.{idx}"  # Use course code + question number
        df_survey.rename(columns={col: course_code_column_name}, inplace=True)

    # Calculate statistics for the survey
    total_students = len(df_survey)
    summary = []

    for question in df_survey.columns[1:]:  # Skip first column (course code column)
        responses = df_survey[question].value_counts().to_dict()
        excellent = responses.get('Excellent', 0)
        very_good = responses.get('Very Good', 0)
        good = responses.get('Good', 0)
        satisfactory = responses.get('Satisfactory', 0)
        poor = responses.get('Poor', 0)
        evg = excellent + very_good + good
        percentage = (evg / total_students) * 100 if total_students > 0 else 0

        summary.append({
            "Subject Code": question,  # Subject code (e.g., 23EVS127.1)
            "Excellent": excellent,
            "Very Good": very_good,
            "Good": good,
            "Satisfactory": satisfactory,
            "Poor": poor,
            "E+V+G": evg,
            "%": round(percentage, 2),  # Add percentage of Excellent + Very Good + Good
        })

    # Generate Word document
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{metadata.get('Subject Code', 'Unknown')}_analysis.docx")
    doc = Document()

    # Add Institute Name and make it centered, bold, and black
    institute_name = "B.N.M. Institute of Technology, Bengaluru-70"
    para = doc.add_paragraph(institute_name, style="Heading 1")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

    # Add Department name (centered, bold, black)
    department_name = "Department of Chemistry"
    para = doc.add_paragraph(department_name, style="Heading 1")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

    # Add "Analysis of course exit survey" (centered, bold, black)
    analysis_text = "Analysis of course exit survey"
    para = doc.add_paragraph(analysis_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

    # Add course information in one line (centered, bold, black)
    course_info = f"Course Code: {metadata.get('Subject Code')}    Course: {metadata.get('Subject Name')}    Branch: {metadata.get('Branch')}    Year: {metadata.get('Year')}"
    para = doc.add_paragraph(course_info)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black

    # Add total students and make it centered, bold, and black
    total_students_text = f"TOTAL NO OF STUDENTS TAKEN SURVEY= {total_students}"
    para = doc.add_paragraph(total_students_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].bold = True
    para.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
    doc.add_paragraph()

    # Add table with responses
    table = doc.add_table(rows=1, cols=8)  # Added one more column for Percentage
    table.style = 'Table Grid'

    # Table headers - set the first column as the course code from metadata
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = metadata.get('Subject Code', 'Unknown')  # Use the actual subject code from metadata
    hdr_cells[1].text = 'EXCELLENT'
    hdr_cells[2].text = 'VERY GOOD'
    hdr_cells[3].text = 'GOOD'
    hdr_cells[4].text = 'SATISFACTORY'
    hdr_cells[5].text = 'POOR'
    hdr_cells[6].text = 'E+V+G'
    hdr_cells[7].text = '%'  # Add a new column header for percentage

    # Fill table with data
    for row in summary:
        row_cells = table.add_row().cells
        row_cells[0].text = row["Subject Code"]  # Subject code (e.g., 23EVS127.1)
        row_cells[1].text = str(row["Excellent"])
        row_cells[2].text = str(row["Very Good"])
        row_cells[3].text = str(row["Good"])
        row_cells[4].text = str(row["Satisfactory"])
        row_cells[5].text = str(row["Poor"])
        row_cells[6].text = str(row["E+V+G"])
        row_cells[7].text = str(row["%"])  # Add percentage value

    # Save the document
    doc.save(output_path)
    return output_path

# Extract metadata from the Excel file
def extract_metadata(df):
    metadata = {}

    # Search for common metadata fields
    for i, row in df.iterrows():
        for val in row:
            if isinstance(val, str):
                if 'Subject Name' in val:
                    metadata['Subject Name'] = val.split(":")[-1].strip()
                elif 'Subject Code' in val:
                    metadata['Subject Code'] = val.split(":")[-1].strip()
                elif 'Branch' in val:
                    metadata['Branch'] = val.split(":")[-1].strip()
                elif 'Year' in val:
                    metadata['Year'] = val.split(":")[-1].strip()

    return metadata

# Find the starting row of survey data based on the first question
def find_question_start(df):
    for i, row in df.iterrows():
        if 'Question' in str(row[0]):
            return i
    return 7  # Default start if nothing is found

if __name__ == '__main__':
    app.run(debug=True)
